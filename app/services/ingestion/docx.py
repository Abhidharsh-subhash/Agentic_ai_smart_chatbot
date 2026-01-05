# app/services/document_processors/docx_processor.py

import re
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from datetime import datetime
from typing import List, Dict, Any, Optional

from app.core.config import settings

# Try to import python-docx
try:
    from docx import Document as DocxDocument

    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    from langchain_community.document_loaders import Docx2txtLoader


def process_docx(
    file_path: str,
    file_metadata: Dict[str, Any],
    chunk_size: int = None,
    chunk_overlap: int = None,
) -> List[Document]:
    """
    Process DOCX file with generalized chunking.
    Uses larger chunks with overlap to preserve context.
    """
    chunk_size = chunk_size or settings.chunk_size
    chunk_overlap = chunk_overlap or settings.chunk_overlap
    filename = file_metadata.get("original_filename", "document")

    try:
        # Extract full text from document
        if HAS_PYTHON_DOCX:
            full_text = _extract_text_with_python_docx(file_path)
        else:
            full_text = _extract_text_with_docx2txt(file_path)

        if not full_text or not full_text.strip():
            return []

        # Create chunks with context
        documents = _create_contextual_chunks(
            text=full_text,
            filename=filename,
            file_metadata=file_metadata,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

        print(
            f"[DOCX] Created {len(documents)} chunks from {filename} "
            f"(chunk_size={chunk_size}, overlap={chunk_overlap})"
        )

        return documents

    except Exception as e:
        raise RuntimeError(f"Error processing DOCX {file_path}: {str(e)}")


def _extract_text_with_python_docx(file_path: str) -> str:
    """Extract text using python-docx, preserving paragraph structure."""
    doc = DocxDocument(file_path)

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                paragraphs.append(" | ".join(row_text))

    return "\n\n".join(paragraphs)


def _extract_text_with_docx2txt(file_path: str) -> str:
    """Extract text using docx2txt as fallback."""
    loader = Docx2txtLoader(file_path)
    documents = loader.load()

    if documents:
        return documents[0].page_content
    return ""


def _create_contextual_chunks(
    text: str,
    filename: str,
    file_metadata: Dict[str, Any],
    chunk_size: int,
    chunk_overlap: int,
) -> List[Document]:
    """
    Create chunks with contextual information.
    Each chunk includes the document name and maintains semantic boundaries.
    """

    # Use semantic separators - try to split at natural boundaries
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=[
            "\n\n\n",  # Multiple blank lines (section breaks)
            "\n\n",  # Paragraph breaks
            "\n",  # Line breaks
            ". ",  # Sentences
            "? ",  # Questions
            "! ",  # Exclamations
            "; ",  # Semicolons
            ", ",  # Commas
            " ",  # Words
            "",  # Characters
        ],
    )

    chunks = splitter.split_text(text)
    documents = []

    for idx, chunk in enumerate(chunks):
        # Add document context to help with retrieval
        # This helps the LLM understand where the content comes from
        context_header = f"Document: {filename}\n\n"

        # For continuation chunks, add a note
        if idx > 0:
            context_header = (
                f"Document: {filename} (part {idx + 1} of {len(chunks)})\n\n"
            )

        contextual_content = context_header + chunk

        doc = Document(
            page_content=contextual_content,
            metadata={
                "file_id": file_metadata.get("file_id"),
                "admin_id": file_metadata.get("admin_id"),
                "folder_id": file_metadata.get("folder_id"),
                "original_filename": file_metadata.get("original_filename"),
                "unique_name": file_metadata.get("unique_name"),
                "file_type": "docx",
                "document_type": "chunk",
                "chunk_index": idx,
                "total_chunks": len(chunks),
                "chunk_size_used": chunk_size,
                "indexed_at": datetime.utcnow().isoformat(),
            },
        )
        documents.append(doc)

    return documents
